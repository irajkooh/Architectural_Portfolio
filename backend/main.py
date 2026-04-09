import os
import threading
from pathlib import Path as _Path
# Load .env from repo root if present (local development)
_env_file = _Path(__file__).parent.parent / ".env"
if _env_file.exists():
    for _line in _env_file.read_text().splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ.setdefault(_k.strip(), _v.strip())
import json
import uuid
import shutil
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, status, Query
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
import secrets
import aiofiles
try:
    import chromadb
    import ollama as ollama_lib
    OLLAMA_DEPS = True
except ImportError:
    OLLAMA_DEPS = False

try:
    import groq as groq_lib
    GROQ_AVAILABLE = True
except Exception:
    GROQ_AVAILABLE = False

try:
    from huggingface_hub import InferenceClient as _HFInferenceClient
    HF_AVAILABLE = True
except Exception:
    HF_AVAILABLE = False

# Chat is available if Groq API key is set, HF_TOKEN is set, or Ollama is installed
def _has_llm() -> bool:
    return (bool(os.environ.get("GROQ_API_KEY"))
            or bool(os.environ.get("HF_TOKEN"))
            or OLLAMA_DEPS)

CHAT_DEPS = True  # kept for admin status endpoint compatibility

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
DATA_DIR = Path(os.environ.get("DATA_DIR", BASE_DIR / "uploads"))
CONFIG_FILE = DATA_DIR / "config.json"
DEFAULTS_FILE = BASE_DIR / "config_defaults.json"
STATIC_DIR = BASE_DIR / "static"          # built React app lands here

ADMIN_PASSWORD    = os.environ.get("ADMIN_PASSWORD", "TAA1346")
SMTP_FILE         = DATA_DIR / "smtp.json"
CHAT_SETTINGS_FILE = DATA_DIR / "chat_settings.json"
CHROMA_DIR        = DATA_DIR / "chroma"
OUTPUT_DIR        = DATA_DIR / "output"
MEDIA_DIR         = BASE_DIR.parent / "media"

def load_chat_settings() -> dict:
    if CHAT_SETTINGS_FILE.exists():
        with open(CHAT_SETTINGS_FILE) as f:
            data = json.load(f)
        if "sample_questions" not in data:
            data["sample_questions"] = []
        return data
    return {
        "ollama_host":      os.environ.get("OLLAMA_HOST", "http://localhost:11434"),
        "chat_model":       os.environ.get("CHAT_MODEL",  "llama3.2"),
        "embed_model":      os.environ.get("EMBED_MODEL", "nomic-embed-text"),
        "sample_questions": [],
    }

def save_chat_settings(s: dict):
    with open(CHAT_SETTINGS_FILE, "w") as f:
        json.dump(s, f, indent=2)

def _chat_cfg():
    s = load_chat_settings()
    return (
        os.environ.get("OLLAMA_HOST")  or s["ollama_host"],
        os.environ.get("CHAT_MODEL")   or s["chat_model"],
        os.environ.get("EMBED_MODEL")  or s["embed_model"],
    )

OLLAMA_HOST  = os.environ.get("OLLAMA_HOST",  "http://localhost:11434")
CHAT_MODEL   = os.environ.get("CHAT_MODEL",   "llama3.2")
EMBED_MODEL  = os.environ.get("EMBED_MODEL",  "nomic-embed-text")

# Ensure dirs exist
for d in ["profile", "resume", "backgrounds", "projects", "showcase_images", "showcase_music", "portfolio"]:
    (DATA_DIR / d).mkdir(parents=True, exist_ok=True)

SHOWCASE_IMAGES_DIR = DATA_DIR / "showcase_images"
SHOWCASE_SLIDESHOW  = SHOWCASE_IMAGES_DIR / "slideshow.mp4"
SHOWCASE_MUSIC_DIR  = DATA_DIR / "showcase_music"

def load_smtp() -> dict:
    if SMTP_FILE.exists():
        with open(SMTP_FILE) as f:
            return json.load(f)
    return {"host": "smtp.gmail.com", "port": 587, "user": "", "pass": ""}

def save_smtp(cfg: dict):
    with open(SMTP_FILE, "w") as f:
        json.dump(cfg, f, indent=2)

# ── Config helpers ─────────────────────────────────────────────────────────────
INFO_FILE = DATA_DIR / "info.xlsx"

def _deep_merge(base: dict, override: dict) -> dict:
    """Fill missing keys in override from base (non-destructive)."""
    result = dict(base)
    for k, v in override.items():
        if k in result and isinstance(result[k], dict) and isinstance(v, dict):
            result[k] = _deep_merge(result[k], v)
        else:
            result[k] = v
    return result

# Fields that need special type casting when read back from Excel strings
_BOOL_FIELDS  = {("hero", "showVideo"), ("resume", "visible")}
_FLOAT_FIELDS = {
    ("theme", "backgroundImageOpacity"),
    ("hero", "backgroundImageOpacity"),
    ("hero", "backgroundVideoOpacity"),
    ("showcase", "videoOpacity"),
}
_NULL_FIELDS  = {
    ("theme", "backgroundImage"), ("theme", "backgroundVideo"),
    ("hero", "backgroundVideo"), ("hero", "backgroundImage"),
    ("about", "photo"), ("resume", "file"), ("portfolio", "file"),
}

def _cast(section: str, key: str, raw) -> object:
    """Cast an Excel cell value back to its correct Python type."""
    if raw is None or raw == "":
        return None if (section, key) in _NULL_FIELDS else ""
    raw = str(raw)
    if (section, key) in _BOOL_FIELDS:
        return raw.lower() == "true"
    if (section, key) in _FLOAT_FIELDS:
        try:
            return float(raw)
        except ValueError:
            return raw
    return raw

def save_to_excel(config: dict):
    """Mirror the full config to info.xlsx every time it changes."""
    try:
        from openpyxl import Workbook
        wb = Workbook()

        # ── Sheet: General (flat section/key/value rows) ──────────────────────
        ws = wb.active
        ws.title = "General"
        ws.append(["Section", "Key", "Value"])

        def _r(section, key, val):
            ws.append([section, key, "" if val is None else str(val)])

        for k, v in config.get("theme", {}).items():
            _r("theme", k, v)

        for k, v in config.get("hero", {}).items():
            _r("hero", k, v)

        about = config.get("about", {})
        for k in ("bio", "philosophy", "linkedin", "photo"):
            _r("about", k, about.get(k))

        for k, v in config.get("contact", {}).items():
            _r("contact", k, v)

        showcase = config.get("showcase", {})
        _r("showcase", "title",    showcase.get("title"))
        _r("showcase", "subtitle", showcase.get("subtitle"))

        resume = config.get("resume", {})
        _r("resume", "visible", resume.get("visible", True))
        _r("resume", "file",    resume.get("file"))

        # ── Sheet: Education ──────────────────────────────────────────────────
        ws2 = wb.create_sheet("Education")
        ws2.append(["Year", "Degree", "Institution"])
        for item in about.get("education", []):
            ws2.append([item.get("year"), item.get("degree"), item.get("institution")])

        # ── Sheet: Skills ─────────────────────────────────────────────────────
        ws3 = wb.create_sheet("Skills")
        ws3.append(["Skill"])
        for skill in about.get("skills", []):
            ws3.append([skill])

        # ── Sheet: Honors ─────────────────────────────────────────────────────
        ws4 = wb.create_sheet("Honors")
        ws4.append(["Year", "Title", "Description", "URL"])
        for item in about.get("honors", []):
            ws4.append([item.get("year"), item.get("title"), item.get("description"), item.get("url")])

        # ── Sheet: Publications ───────────────────────────────────────────────
        ws5 = wb.create_sheet("Publications")
        ws5.append(["Year", "Title", "Venue", "URL"])
        for item in about.get("publications", []):
            ws5.append([item.get("year"), item.get("title"), item.get("venue"), item.get("url")])

        # ── Sheet: Projects ───────────────────────────────────────────────────
        ws6 = wb.create_sheet("Projects")
        ws6.append(["ID", "Title", "Name", "Category", "Year", "Client", "Location", "Description", "Cover", "Video", "VideoURL", "Images"])
        for proj in config.get("projects", []):
            ws6.append([
                proj.get("id", ""),
                proj.get("title", ""),
                proj.get("name", "") or "",
                proj.get("category", ""),
                proj.get("year", ""),
                proj.get("client", ""),
                proj.get("location", ""),
                proj.get("description", ""),
                proj.get("cover") or "",
                proj.get("video") or "",
                proj.get("videoUrl") or "",
                "|".join(proj.get("images") or []),
            ])

        # ── Sheet: ProjectTypes ───────────────────────────────────────────────
        ws7 = wb.create_sheet("ProjectTypes")
        ws7.append(["Type"])
        for t in config.get("projectTypes", []):
            ws7.append([t])

        # ── Sheet: Portfolio ──────────────────────────────────────────────────
        portfolio = config.get("portfolio", {})
        _r("portfolio", "visible", portfolio.get("visible", True))
        _r("portfolio", "file",    portfolio.get("file"))

        # ── Sheet: Resume_Highlights ──────────────────────────────────────────
        ws8 = wb.create_sheet("Resume_Highlights")
        ws8.append(["Highlight"])
        for h in resume.get("highlights", []):
            ws8.append([h])

        # ── Sheet: Showcase_Videos ────────────────────────────────────────────
        ws9 = wb.create_sheet("Showcase_Videos")
        ws9.append(["ID", "URL", "Title", "Description"])
        for v in showcase.get("videos", []):
            ws9.append([
                v.get("id", ""),
                v.get("url") or "",
                v.get("title") or "",
                v.get("description") or "",
            ])

        wb.save(INFO_FILE)
        print("[info.xlsx] Saved.")
    except Exception as e:
        print(f"[info.xlsx] Save failed: {e}")


def load_from_excel() -> dict:
    """Reconstruct the full config dict from info.xlsx."""
    from openpyxl import load_workbook
    wb = load_workbook(INFO_FILE, data_only=True)

    config: dict = {
        "theme": {}, "hero": {},
        "about": {"education": [], "skills": [], "honors": [], "publications": []},
        "contact": {}, "showcase": {"videos": []},
        "projects": [], "projectTypes": [],
        "resume": {"highlights": []},
        "portfolio": {},
    }

    # General
    if "General" in wb.sheetnames:
        for row in wb["General"].iter_rows(min_row=2, values_only=True):
            sec, key, raw = (row[0] or ""), (row[1] or ""), row[2]
            if not sec or not key:
                continue
            val = _cast(sec, key, raw)
            if sec in ("theme", "hero", "contact"):
                config[sec][key] = val
            elif sec == "about":
                config["about"][key] = val
            elif sec == "showcase":
                config["showcase"][key] = val
            elif sec == "resume":
                config["resume"][key] = val
            elif sec == "portfolio":
                config["portfolio"][key] = val

    # Education
    if "Education" in wb.sheetnames:
        for row in wb["Education"].iter_rows(min_row=2, values_only=True):
            if any(c for c in row if c is not None and c != ""):
                config["about"]["education"].append({
                    "year": str(row[0] or ""), "degree": str(row[1] or ""),
                    "institution": str(row[2] or ""),
                })

    # Skills
    if "Skills" in wb.sheetnames:
        for row in wb["Skills"].iter_rows(min_row=2, values_only=True):
            if row[0]:
                config["about"]["skills"].append(str(row[0]))

    # Honors
    if "Honors" in wb.sheetnames:
        for row in wb["Honors"].iter_rows(min_row=2, values_only=True):
            if any(c for c in row if c is not None and c != ""):
                config["about"]["honors"].append({
                    "year": str(row[0] or ""), "title": str(row[1] or ""),
                    "description": str(row[2] or ""), "url": str(row[3] or ""),
                })

    # Publications
    if "Publications" in wb.sheetnames:
        for row in wb["Publications"].iter_rows(min_row=2, values_only=True):
            if any(c for c in row if c is not None and c != ""):
                config["about"]["publications"].append({
                    "year": str(row[0] or ""), "title": str(row[1] or ""),
                    "venue": str(row[2] or ""), "url": str(row[3] or ""),
                })

    # Projects
    if "Projects" in wb.sheetnames:
        xlsx_header = None
        for row in wb["Projects"].iter_rows(min_row=1, max_row=1, values_only=True):
            xlsx_header = [str(c or "").strip().lower() for c in row]
        is_new_fmt = xlsx_header is not None and len(xlsx_header) >= 12 and "name" in xlsx_header[2:3]
        for row in wb["Projects"].iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            # Detect format by header, not by per-row content (avoids empty-name bug)
            if is_new_fmt:
                # New format: ID, Title, Name, Category, Year, Client, Location, Desc, Cover, Video, VideoURL, Images
                name_val = str(row[2] or "").strip()
                cat_col = 3
            else:
                # Old format: ID, Title, Category, Year, Client, Location, Desc, Cover, Video, VideoURL, Images
                name_val = ""
                cat_col = 2
            config["projects"].append({
                "id":          str(row[0]),
                "title":       str(row[1] or ""),
                "name":        name_val,
                "category":    str(row[cat_col] or ""),
                "year":        str(row[cat_col+1] or ""),
                "client":      str(row[cat_col+2] or ""),
                "location":    str(row[cat_col+3] or ""),
                "description": str(row[cat_col+4] or ""),
                "cover":       str(row[cat_col+5] or "") or None,
                "video":       str(row[cat_col+6] or "") or None,
                "videoUrl":    str(row[cat_col+7] or "") or None,
                "images":      [i for i in str(row[cat_col+8] or "").split("|") if i],
            })

    # ProjectTypes
    if "ProjectTypes" in wb.sheetnames:
        for row in wb["ProjectTypes"].iter_rows(min_row=2, values_only=True):
            if row[0]:
                config["projectTypes"].append(str(row[0]))

    # Resume_Highlights
    if "Resume_Highlights" in wb.sheetnames:
        for row in wb["Resume_Highlights"].iter_rows(min_row=2, values_only=True):
            if row[0]:
                config["resume"]["highlights"].append(str(row[0]))

    # Showcase_Videos
    if "Showcase_Videos" in wb.sheetnames:
        for row in wb["Showcase_Videos"].iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            config["showcase"]["videos"].append({
                "id":          str(row[0]),
                "url":         str(row[1] or "") or None,
                "title":       str(row[2] or ""),
                "description": str(row[3] or ""),
            })

    return config


def _empty_config() -> dict:
    """Clean slate — theme defaults kept, all personal content empty."""
    with open(DEFAULTS_FILE) as f:
        defaults = json.load(f)
    return {
        "theme": defaults["theme"],
        "hero": {
            "title": "", "subtitle": "", "tagline": "",
            "backgroundVideo": None, "backgroundImage": None,
            "showVideo": False,
            "ctaPortfolio": "View Portfolio", "ctaResume": "Download Resume",
        },
        "about": {
            "bio": "", "philosophy": "", "linkedin": "", "photo": None,
            "education": [], "skills": [], "honors": [], "publications": [],
        },
        "projectTypes": defaults["projectTypes"],
        "projects": [],
        "resume": {"file": None, "visible": True, "highlights": []},
        "portfolio": {"file": None, "visible": True},
        "contact": {"email": "", "linkedin": "", "message": ""},
        "showcase": {"title": "Showcase", "subtitle": "", "videos": [], "videoOpacity": 1.0},
    }


def _normalize_projects(projects: list) -> list:
    """Ensure every project dict has all expected fields with safe defaults.
    This handles configs saved before a field was introduced (e.g. 'name' added
    after initial deployment) so the UI always receives a complete object."""
    for p in projects:
        p.setdefault("name", "")
        p.setdefault("videoUrl", None)
        p.setdefault("video", None)
        p.setdefault("cover", None)
        p.setdefault("images", [])
        p.setdefault("client", "")
        p.setdefault("location", "")
        p.setdefault("year", "")
    return projects


def _sync_file_paths(cfg: dict) -> None:
    """Ensure portfolio/resume file fields match what is actually on disk.
    Prevents config drift where file=null but the PDF was successfully uploaded.
    Called on every load_config() so any subsequent save_config() persists the
    correct path — matching how resume.pdf (always in the Docker image) works."""
    if (DATA_DIR / "portfolio" / "portfolio.pdf").exists():
        if not (cfg.get("portfolio") or {}).get("file"):
            cfg.setdefault("portfolio", {})["file"] = "/uploads/portfolio/portfolio.pdf"
    if (DATA_DIR / "resume" / "resume.pdf").exists():
        if not (cfg.get("resume") or {}).get("file"):
            cfg.setdefault("resume", {})["file"] = "/uploads/resume/resume.pdf"


def load_config() -> dict:
    with open(DEFAULTS_FILE) as f:
        defaults = json.load(f)

    # Priority 1: config.json — active session state, most up to date
    if CONFIG_FILE.exists():
        with open(CONFIG_FILE) as f:
            saved = json.load(f)
        merged = _deep_merge(defaults, saved)
        if not merged.get("projects") and defaults.get("projects"):
            merged["projects"] = defaults["projects"]
        merged["projects"] = _normalize_projects(merged.get("projects", []))
        _sync_file_paths(merged)  # recover file paths nulled by HF config drift
        return merged

    # Priority 2: info.xlsx — portable deployment snapshot
    if INFO_FILE.exists():
        try:
            config = load_from_excel()
            config["projects"] = _normalize_projects(config.get("projects", []))
            _sync_file_paths(config)  # recover file paths nulled by HF config drift
            # Persist to config.json so subsequent requests are fast
            with open(CONFIG_FILE, "w") as f:
                json.dump(config, f, indent=2)
            print("[info.xlsx] Bootstrapped config.json from info.xlsx")
            return config
        except Exception as e:
            print(f"[info.xlsx] Load failed, starting empty: {e}")

    # Priority 3: no files at all — clean empty start
    return _empty_config()


def save_config(cfg: dict):
    with open(CONFIG_FILE, "w") as f:
        json.dump(cfg, f, indent=2)
    save_to_excel(cfg)  # keep info.xlsx in sync
    # Push immediately in background thread so it never blocks the request
    threading.Thread(target=_hf_push_file, args=(CONFIG_FILE,), daemon=True).start()

# ── App ────────────────────────────────────────────────────────────────────────
from contextlib import asynccontextmanager
from apscheduler.schedulers.asyncio import AsyncIOScheduler
import datetime

HF_DATASET_REPO = os.environ.get("HF_DATASET_REPO", "")
HF_TOKEN        = os.environ.get("HF_TOKEN", "")

# Files/dirs to never push to the dataset (large or ephemeral)
_SYNC_SKIP = ("chroma", ".DS_Store", "smtp.json", "output")

def _hf_pull():
    """Pull all files from the HF dataset repo into DATA_DIR on startup."""
    if not HF_DATASET_REPO or not HF_TOKEN:
        return
    try:
        from huggingface_hub import snapshot_download
        snapshot_download(
            repo_id=HF_DATASET_REPO,
            repo_type="dataset",
            token=HF_TOKEN,
            local_dir=str(DATA_DIR),
            ignore_patterns=["*.gitattributes", ".gitattributes", "README.md"],
        )
        print(f"[dataset] Pulled data from {HF_DATASET_REPO}")
    except Exception as e:
        print(f"[dataset] Pull failed, using local data: {e}")

_sync_mtimes: dict = {}

def _hf_push_file(path: Path):
    """Immediately push a single file to the HF dataset repo."""
    if not HF_DATASET_REPO or not HF_TOKEN:
        return
    try:
        from huggingface_hub import upload_file
        rel = str(path.relative_to(DATA_DIR))
        upload_file(
            path_or_fileobj=str(path),
            path_in_repo=rel,
            repo_id=HF_DATASET_REPO,
            repo_type="dataset",
            token=HF_TOKEN,
        )
        _sync_mtimes[str(path)] = path.stat().st_mtime
        print(f"[dataset] Pushed {rel} immediately")
    except Exception as e:
        print(f"[dataset] Immediate push failed for {path.name}: {e}")

def _hf_push_changed():
    """Push any locally modified files to the HF dataset repo."""
    if not HF_DATASET_REPO or not HF_TOKEN:
        return
    try:
        from huggingface_hub import upload_file
        for f in DATA_DIR.rglob("*"):
            if not f.is_file():
                continue
            if any(skip in str(f) for skip in _SYNC_SKIP):
                continue
            mtime = f.stat().st_mtime
            if _sync_mtimes.get(str(f)) == mtime:
                continue
            _sync_mtimes[str(f)] = mtime
            rel = str(f.relative_to(DATA_DIR))
            try:
                upload_file(
                    path_or_fileobj=str(f),
                    path_in_repo=rel,
                    repo_id=HF_DATASET_REPO,
                    repo_type="dataset",
                    token=HF_TOKEN,
                )
                print(f"[dataset] Pushed {rel}")
            except Exception as e:
                print(f"[dataset] Failed to push {rel}: {e}")
    except Exception as e:
        print(f"[dataset] Sync error: {e}")

def daily_keepalive():
    print(f"[scheduler] daily keepalive — {datetime.datetime.utcnow().isoformat()}Z")
    save_to_excel(load_config())

def _bootstrap_smtp():
    """Write smtp.json from env vars on first startup (e.g. HF Spaces secrets)."""
    host = os.environ.get("SMTP_HOST", "")
    port = os.environ.get("SMTP_PORT", "")
    user = os.environ.get("SMTP_USER", "")
    pw   = os.environ.get("SMTP_PASS", "")
    if not (user and pw):
        return
    existing = load_smtp()
    if existing.get("user") == user and existing.get("pass") == pw:
        return  # already up to date
    save_smtp({
        "host": host or existing.get("host", "smtp.gmail.com"),
        "port": int(port) if port else existing.get("port", 587),
        "user": user,
        "pass": pw,
    })
    print("[smtp] Bootstrapped smtp.json from environment variables.")

@asynccontextmanager
async def lifespan(app):
    # Save Docker baked-in config BEFORE _hf_pull() may overwrite it.
    # Used to migrate fields (like 'name') that exist in the deployed codebase
    # but may be absent in an older HF dataset snapshot.
    _seed_cfg: dict | None = None
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE) as _f:
                _seed_cfg = json.load(_f)
        except Exception:
            pass

    _hf_pull()

    # Migration: copy missing project fields from seed (Docker) into live (HF-pulled) config.
    # Specifically recovers 'name' (and any future fields) if the HF dataset has older data.
    if _seed_cfg and CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE) as _f:
                _live = json.load(_f)
            _seed_map = {p["id"]: p for p in _seed_cfg.get("projects", []) if p.get("id")}
            _changed = False
            for _proj in _live.get("projects", []):
                _seed = _seed_map.get(_proj.get("id", ""))
                if _seed:
                    # Copy any field that is missing or empty in live but present in seed
                    for _field in ("name", "client", "location", "year"):
                        if not _proj.get(_field) and _seed.get(_field):
                            _proj[_field] = _seed[_field]
                            _changed = True
            if _changed:
                with open(CONFIG_FILE, "w") as _f:
                    json.dump(_live, _f, indent=2)
                print("[migrate] Recovered project fields from deployment config into HF-pulled config")
        except Exception as _e:
            print(f"[migrate] Field migration failed: {_e}")

    # Startup file-path recovery: if PDFs exist on disk (either from Docker image or HF pull)
    # but config.json has their paths as null, restore them.  This is the permanent fix for
    # the portfolio download button disappearing on Space after any config save.
    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE) as _f:
                _live_fp = json.load(_f)
            _fp_changed = False
            if (DATA_DIR / "portfolio" / "portfolio.pdf").exists():
                if not (_live_fp.get("portfolio") or {}).get("file"):
                    _live_fp.setdefault("portfolio", {})["file"] = "/uploads/portfolio/portfolio.pdf"
                    _fp_changed = True
            if (DATA_DIR / "resume" / "resume.pdf").exists():
                if not (_live_fp.get("resume") or {}).get("file"):
                    _live_fp.setdefault("resume", {})["file"] = "/uploads/resume/resume.pdf"
                    _fp_changed = True
            if _fp_changed:
                with open(CONFIG_FILE, "w") as _f:
                    json.dump(_live_fp, _f, indent=2)
                print("[migrate] Restored PDF file paths in config.json from disk")
        except Exception as _e:
            print(f"[migrate] File-path recovery failed: {_e}")

    _bootstrap_smtp()
    # Only sync Excel if we actually have config data; prevents overwriting with empty on pull failure
    if CONFIG_FILE.exists():
        save_to_excel(load_config())

    scheduler = AsyncIOScheduler()
    scheduler.add_job(daily_keepalive, "cron", hour=6, minute=0)
    scheduler.add_job(_hf_push_changed, "interval", minutes=10)
    scheduler.start()

    yield

    _hf_push_changed()   # final push on shutdown
    scheduler.shutdown()

app = FastAPI(title="Tella Portfolio API", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

security = HTTPBasic()

def verify_admin(credentials: HTTPBasicCredentials = Depends(security)):
    correct = secrets.compare_digest(credentials.password, ADMIN_PASSWORD)
    if not correct:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid admin password",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username

# ── Public API ─────────────────────────────────────────────────────────────────
@app.get("/api/config")
def get_config():
    return load_config()

# ── Admin Auth check ───────────────────────────────────────────────────────────
@app.get("/api/admin/verify")
def verify(_: str = Depends(verify_admin)):
    return {"ok": True}

# ── Info file export / import ──────────────────────────────────────────────────
@app.get("/api/admin/info-export")
def export_info(_: str = Depends(verify_admin)):
    """Download info.xlsx — use this to back up or migrate your content."""
    if not INFO_FILE.exists():
        raise HTTPException(status_code=404, detail="No info.xlsx found yet. Save some content first.")
    return FileResponse(
        INFO_FILE,
        filename="info.xlsx",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

@app.post("/api/admin/info-import")
async def import_info(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    """Upload a previously exported info.xlsx to restore all content."""
    data = await file.read()
    INFO_FILE.write_bytes(data)
    try:
        config = load_from_excel()
        save_config(config)
        return {"ok": True, "message": "Config restored from info.xlsx"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not parse info.xlsx: {e}")

# ── Theme ──────────────────────────────────────────────────────────────────────
@app.post("/api/admin/theme")
async def update_theme(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["theme"].update(data)
    save_config(cfg)
    return cfg["theme"]

@app.post("/api/admin/theme/background-image")
async def upload_bg_image(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "backgrounds" / f"background{ext}"
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    cfg = load_config()
    cfg["theme"]["backgroundImage"] = f"/uploads/backgrounds/background{ext}"
    save_config(cfg)
    return {"backgroundImage": cfg["theme"]["backgroundImage"]}

@app.delete("/api/admin/theme/background-image")
def delete_bg_image(_: str = Depends(verify_admin)):
    cfg = load_config()
    if cfg["theme"].get("backgroundImage"):
        p = DATA_DIR / "backgrounds"
        for f in p.iterdir():
            f.unlink(missing_ok=True)
        cfg["theme"]["backgroundImage"] = None
        save_config(cfg)
    return {"ok": True}

# ── Hero ───────────────────────────────────────────────────────────────────────
@app.post("/api/admin/hero")
async def update_hero(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["hero"].update(data)
    save_config(cfg)
    return cfg["hero"]

@app.post("/api/admin/hero/video")
async def upload_hero_video(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "backgrounds" / f"hero_video{ext}"
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    url = f"/uploads/backgrounds/hero_video{ext}"
    cfg = load_config()
    cfg["hero"]["backgroundVideo"] = url
    cfg["hero"]["showVideo"] = True
    save_config(cfg)
    return {"backgroundVideo": url}

@app.delete("/api/admin/hero/video")
def delete_hero_video(_: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["hero"]["backgroundVideo"] = None
    cfg["hero"]["showVideo"] = False
    save_config(cfg)
    return {"ok": True}

@app.post("/api/admin/hero/image")
async def upload_hero_image(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "backgrounds" / f"hero_image{ext}"
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    url = f"/uploads/backgrounds/hero_image{ext}"
    cfg = load_config()
    cfg["hero"]["backgroundImage"] = url
    save_config(cfg)
    return {"backgroundImage": url}

@app.delete("/api/admin/hero/image")
def delete_hero_image(_: str = Depends(verify_admin)):
    cfg = load_config()
    if cfg["hero"].get("backgroundImage"):
        p = DATA_DIR / "backgrounds"
        for f in p.glob("hero_image.*"):
            f.unlink(missing_ok=True)
        cfg["hero"]["backgroundImage"] = None
        save_config(cfg)
    return {"ok": True}

# ── About ──────────────────────────────────────────────────────────────────────
@app.post("/api/admin/about")
async def update_about(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["about"].update(data)
    save_config(cfg)
    return cfg["about"]

@app.post("/api/admin/about/photo")
async def upload_photo(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "profile" / f"photo{ext}"
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    url = f"/uploads/profile/photo{ext}"
    cfg = load_config()
    cfg["about"]["photo"] = url
    save_config(cfg)
    return {"photo": url}

@app.delete("/api/admin/about/photo")
def delete_photo(_: str = Depends(verify_admin)):
    cfg = load_config()
    for f in (DATA_DIR / "profile").iterdir():
        f.unlink(missing_ok=True)
    cfg["about"]["photo"] = None
    save_config(cfg)
    return {"ok": True}

# ── Projects ───────────────────────────────────────────────────────────────────
@app.post("/api/admin/projects")
async def create_project(
    title: str = Form(...),
    name: str = Form(""),
    description: str = Form(...),
    category: str = Form(...),
    year: str = Form(""),
    client: str = Form(""),
    location: str = Form(""),
    _: str = Depends(verify_admin),
):
    project_id = str(uuid.uuid4())
    (DATA_DIR / "projects" / project_id / "images").mkdir(parents=True, exist_ok=True)
    project = {
        "id": project_id,
        "title": title,
        "name": name,
        "description": description,
        "category": category,
        "year": year,
        "client": client,
        "location": location,
        "images": [],
        "cover": None,
        "video": None,
        "videoUrl": None,
    }
    cfg = load_config()
    cfg["projects"].append(project)
    save_config(cfg)
    return project

@app.put("/api/admin/projects/{project_id}")
async def update_project(project_id: str, data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    for i, p in enumerate(cfg["projects"]):
        if p["id"] == project_id:
            cfg["projects"][i].update(data)
            save_config(cfg)
            return cfg["projects"][i]
    raise HTTPException(404, "Project not found")

@app.post("/api/admin/projects/reorder")
async def reorder_projects(data: dict, _: str = Depends(verify_admin)):
    """Reorder projects to match the supplied list of IDs."""
    ids: list[str] = data.get("ids", [])
    cfg = load_config()
    lookup = {p["id"]: p for p in cfg["projects"]}
    reordered = [lookup[i] for i in ids if i in lookup]
    # append any projects not mentioned (safety)
    mentioned = set(ids)
    for p in cfg["projects"]:
        if p["id"] not in mentioned:
            reordered.append(p)
    cfg["projects"] = reordered
    save_config(cfg)
    return {"ok": True}

@app.delete("/api/admin/projects/{project_id}")
def delete_project(project_id: str, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["projects"] = [p for p in cfg["projects"] if p["id"] != project_id]
    save_config(cfg)
    proj_dir = DATA_DIR / "projects" / project_id
    if proj_dir.exists():
        shutil.rmtree(proj_dir)
    return {"ok": True}

@app.post("/api/admin/projects/{project_id}/images")
async def upload_project_image(
    project_id: str,
    file: UploadFile = File(...),
    _: str = Depends(verify_admin),
):
    img_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "projects" / project_id / "images" / f"{img_id}{ext}"
    dest.parent.mkdir(parents=True, exist_ok=True)
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    url = f"/uploads/projects/{project_id}/images/{img_id}{ext}"
    cfg = load_config()
    for p in cfg["projects"]:
        if p["id"] == project_id:
            p["images"].append(url)
            if not p["cover"]:
                p["cover"] = url
            break
    save_config(cfg)
    return {"url": url}

@app.delete("/api/admin/projects/{project_id}/images")
def delete_project_image(project_id: str, url: str, _: str = Depends(verify_admin)):
    cfg = load_config()
    for p in cfg["projects"]:
        if p["id"] == project_id:
            p["images"] = [i for i in p["images"] if i != url]
            if p["cover"] == url:
                p["cover"] = p["images"][0] if p["images"] else None
            break
    save_config(cfg)
    filename = Path(url).name
    f = DATA_DIR / "projects" / project_id / "images" / filename
    f.unlink(missing_ok=True)
    return {"ok": True}

@app.post("/api/admin/projects/{project_id}/video")
async def upload_project_video(
    project_id: str,
    file: UploadFile = File(...),
    _: str = Depends(verify_admin),
):
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "projects" / project_id / f"video{ext}"
    dest.parent.mkdir(parents=True, exist_ok=True)
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    url = f"/uploads/projects/{project_id}/video{ext}"
    cfg = load_config()
    for p in cfg["projects"]:
        if p["id"] == project_id:
            p["video"] = url
            break
    save_config(cfg)
    return {"video": url}

@app.delete("/api/admin/projects/{project_id}/video")
def delete_project_video(project_id: str, _: str = Depends(verify_admin)):
    cfg = load_config()
    for p in cfg["projects"]:
        if p["id"] == project_id:
            p["video"] = None
            p["videoUrl"] = None
            break
    save_config(cfg)
    return {"ok": True}

# ── Resume ─────────────────────────────────────────────────────────────────────
@app.post("/api/admin/resume")
async def update_resume_meta(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["resume"].update(data)
    save_config(cfg)
    return cfg["resume"]

@app.get("/api/resume/preview")
async def resume_preview():
    pdf_path = DATA_DIR / "resume" / "resume.pdf"
    if not pdf_path.exists():
        raise HTTPException(404, "Resume not found")
    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        headers={
            "Content-Disposition": "inline; filename=resume.pdf",
            "X-Frame-Options": "ALLOWALL",
            "Content-Security-Policy": "frame-ancestors *",
        },
    )

@app.get("/api/portfolio/preview")
async def portfolio_preview():
    pdf_path = DATA_DIR / "portfolio" / "portfolio.pdf"
    if not pdf_path.exists():
        raise HTTPException(404, "Portfolio not found")
    return FileResponse(
        str(pdf_path),
        media_type="application/pdf",
        headers={
            "Content-Disposition": "inline; filename=portfolio.pdf",
            "X-Frame-Options": "ALLOWALL",
            "Content-Security-Policy": "frame-ancestors *",
        },
    )

@app.post("/api/admin/resume/file")
async def upload_resume(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    dest = DATA_DIR / "resume" / "resume.pdf"
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    cfg = load_config()
    cfg["resume"]["file"] = "/uploads/resume/resume.pdf"
    save_config(cfg)
    threading.Thread(target=_hf_push_file, args=(dest,), daemon=True).start()
    ingested = 0
    if CHAT_DEPS:
        try:
            ingested = ingest_resume_to_chroma()
        except Exception:
            pass  # Ollama may not be running — ingest manually from Admin → Chat
    return {"file": cfg["resume"]["file"], "ingested_chunks": ingested}

@app.delete("/api/admin/resume/file")
def delete_resume(_: str = Depends(verify_admin)):
    (DATA_DIR / "resume" / "resume.pdf").unlink(missing_ok=True)
    cfg = load_config()
    cfg["resume"]["file"] = None
    save_config(cfg)
    return {"ok": True}

# ── Portfolio file (uploaded PDF) ──────────────────────────────────────────────
@app.get("/api/admin/portfolio")
def get_portfolio_settings(_: str = Depends(verify_admin)):
    cfg = load_config()
    return cfg.get("portfolio", {"file": None, "visible": True})

@app.post("/api/admin/portfolio")
async def save_portfolio_settings(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    if "portfolio" not in cfg:
        cfg["portfolio"] = {"file": None, "visible": True}
    if "visible" in data:
        cfg["portfolio"]["visible"] = bool(data["visible"])
    if "file" in data:
        cfg["portfolio"]["file"] = data["file"] or None
    save_config(cfg)
    ingested = 0
    if "file" in data and data["file"] and CHAT_DEPS:
        try:
            ingested = ingest_resume_to_chroma()
        except Exception:
            pass
    return {**cfg["portfolio"], "ingested_chunks": ingested}

@app.post("/api/admin/portfolio/file")
async def upload_portfolio(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    (DATA_DIR / "portfolio").mkdir(exist_ok=True)
    dest = DATA_DIR / "portfolio" / "portfolio.pdf"
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    cfg = load_config()
    if "portfolio" not in cfg:
        cfg["portfolio"] = {"visible": True}
    cfg["portfolio"]["file"] = "/uploads/portfolio/portfolio.pdf"
    save_config(cfg)
    threading.Thread(target=_hf_push_file, args=(dest,), daemon=True).start()
    ingested = 0
    if CHAT_DEPS:
        try:
            ingested = ingest_resume_to_chroma()
        except Exception:
            pass
    return {"file": cfg["portfolio"]["file"], "ingested_chunks": ingested}

@app.delete("/api/admin/portfolio/file")
def delete_portfolio(_: str = Depends(verify_admin)):
    (DATA_DIR / "portfolio" / "portfolio.pdf").unlink(missing_ok=True)
    cfg = load_config()
    if "portfolio" in cfg:
        cfg["portfolio"]["file"] = None
    save_config(cfg)
    return {"ok": True}

# ── Project Types ──────────────────────────────────────────────────────────────
@app.post("/api/admin/project-types")
async def update_project_types(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["projectTypes"] = data.get("projectTypes", cfg.get("projectTypes", []))
    save_config(cfg)
    return {"projectTypes": cfg["projectTypes"]}

# ── RAG helpers ────────────────────────────────────────────────────────────────
def _detect_device() -> str:
    import platform
    sys_name = platform.system()
    machine  = platform.machine()
    if sys_name == "Darwin" and machine == "arm64":
        return "mps"
    try:
        import subprocess as _sp
        r = _sp.run(
            ["nvidia-smi", "--query-gpu=name", "--format=csv,noheader"],
            capture_output=True, text=True, timeout=3,
        )
        if r.returncode == 0 and r.stdout.strip():
            return f"cuda ({r.stdout.strip().splitlines()[0]})"
    except Exception:
        pass
    return "cpu"

def _detect_location() -> str:
    space_id = os.environ.get("SPACE_ID")
    if space_id:
        org, name = space_id.split("/") if "/" in space_id else (space_id, space_id)
        return f"https://{org}-{name}.hf.space"
    return "http://localhost:5173"

@app.get("/api/chat/info")
def chat_info():
    s = load_chat_settings()
    if os.environ.get("GROQ_API_KEY"):
        model = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile") + " (Groq)"
    elif os.environ.get("HF_TOKEN"):
        model = os.environ.get("HF_MODEL", "meta-llama/Llama-3.1-8B-Instruct") + " (HuggingFace)"
    else:
        model = os.environ.get("CHAT_MODEL") or s.get("chat_model", "llama3.2")
    return {
        "device":   _detect_device(),
        "model":    model,
        "location": _detect_location(),
    }

def _check_chat_deps():
    if not _has_llm():
        raise HTTPException(503, "No LLM configured. Set GROQ_API_KEY, HF_TOKEN, or install ollama.")


def _llm_chat(messages: list) -> str:
    """Send chat messages to LLM. Priority: Groq → HuggingFace → Ollama."""
    groq_key = os.environ.get("GROQ_API_KEY")
    if groq_key:
        if not GROQ_AVAILABLE:
            raise RuntimeError("groq package not installed")
        groq_model = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
        client = groq_lib.Groq(api_key=groq_key)
        resp = client.chat.completions.create(model=groq_model, messages=messages)
        return resp.choices[0].message.content.strip()
    hf_token = os.environ.get("HF_TOKEN")
    if hf_token:
        if not HF_AVAILABLE:
            raise RuntimeError("huggingface_hub package not installed")
        hf_model = os.environ.get("HF_MODEL", "meta-llama/Llama-3.1-8B-Instruct")
        client = _HFInferenceClient(model=hf_model, token=hf_token)
        resp = client.chat_completion(messages=messages, max_tokens=1024)
        return resp.choices[0].message.content.strip()
    if not OLLAMA_DEPS:
        raise RuntimeError("No LLM available. Set GROQ_API_KEY, HF_TOKEN, or install ollama.")
    ollama_host, chat_model, _ = _chat_cfg()
    client = ollama_lib.Client(host=ollama_host)
    response = client.chat(model=chat_model, messages=messages)
    msg = response.message if hasattr(response, "message") else response["message"]
    return (msg.content if hasattr(msg, "content") else msg["content"]).strip()


def _embed(texts: list) -> list:
    """Embed a list of strings via Ollama (only used for ChromaDB ingest)."""
    if not OLLAMA_DEPS:
        raise RuntimeError("ollama not installed — ChromaDB ingest unavailable")
    host, _, embed_model = _chat_cfg()
    client = ollama_lib.Client(host=host)
    result = []
    for text in texts:
        resp = client.embeddings(model=embed_model, prompt=text)
        emb  = resp.embedding if hasattr(resp, "embedding") else resp["embedding"]
        result.append(emb)
    return result

def _chroma_collection(reset: bool = False):
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(
        path=str(CHROMA_DIR),
        settings=chromadb.Settings(anonymized_telemetry=False),
    )
    if reset:
        try:
            client.delete_collection("resume")
        except Exception:
            pass
    # No embedding function — we pass pre-computed embeddings every time
    return client.get_or_create_collection("resume")

def _extract_pdf_text(pdf_path) -> str:
    """Extract text from PDF, trying pdfplumber first for better layout handling."""
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
                pages.append(text)
        return "\n\n".join(pages)
    except ImportError:
        pass
    import pypdf
    reader = pypdf.PdfReader(str(pdf_path))
    return "\n\n".join(p.extract_text() or "" for p in reader.pages)

def _split_text(text: str, size: int = 1000, overlap: int = 150) -> list:
    """Paragraph-aware chunking: prefer splitting on blank lines."""
    import re
    # Normalize whitespace: collapse 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]

    chunks, current = [], ""
    for para in paragraphs:
        if not current:
            current = para
        elif len(current) + len(para) + 2 <= size:
            current += "\n\n" + para
        else:
            chunks.append(current)
            # overlap: carry last paragraph into next chunk
            current = (para if len(para) <= size
                       else current[-overlap:] + "\n\n" + para)
    if current:
        chunks.append(current)
    return [c for c in chunks if len(c) > 30]

def ingest_resume_to_chroma() -> int:
    """Extract text from resume + portfolio PDFs and store for chatbot context."""
    cfg = load_config()
    texts: list[str] = []

    # Resume PDF
    resume_url = cfg.get("resume", {}).get("file", "")
    if resume_url:
        rp = DATA_DIR / resume_url.replace("/uploads/", "", 1)
        if rp.exists():
            try:
                t = _extract_pdf_text(rp)
                if t.strip():
                    texts.append(t)
            except Exception:
                pass

    # Portfolio PDF (uploaded via admin)
    portfolio_url = cfg.get("portfolio", {}).get("file", "")
    if portfolio_url and portfolio_url.startswith("/uploads/"):
        pp = DATA_DIR / portfolio_url.replace("/uploads/", "", 1)
        if pp.exists():
            try:
                t = _extract_pdf_text(pp)
                if t.strip():
                    texts.append(t)
            except Exception:
                pass

    if not texts:
        return 0

    full_text = "\n\n".join(texts)
    # Always save full text — chat uses this directly
    (DATA_DIR / "resume").mkdir(exist_ok=True)
    (DATA_DIR / "resume" / "resume_text.txt").write_text(full_text, encoding="utf-8")
    # ChromaDB vector store — optional, only when Ollama is available
    if OLLAMA_DEPS:
        try:
            chunks = _split_text(full_text)
            embeddings = _embed(chunks)
            col = _chroma_collection(reset=True)
            col.upsert(documents=chunks, embeddings=embeddings, ids=[f"chunk_{i}" for i in range(len(chunks))])
        except Exception:
            pass  # ChromaDB/Ollama unavailable — text file is enough
    return len(texts)  # number of PDFs ingested

# ── Chat / RAG endpoints ────────────────────────────────────────────────────────
@app.get("/api/admin/chat/settings")
def get_chat_settings(_: str = Depends(verify_admin)):
    return load_chat_settings()

@app.post("/api/admin/chat/settings")
async def update_chat_settings(data: dict, _: str = Depends(verify_admin)):
    s = load_chat_settings()
    for k in ("ollama_host", "chat_model", "embed_model", "sample_questions"):
        if k in data:
            s[k] = data[k]
    save_chat_settings(s)
    return s

@app.get("/api/chat/sample-questions")
def get_sample_questions():
    return load_chat_settings().get("sample_questions", [])

@app.get("/api/admin/chat/status")
def chat_status(_: str = Depends(verify_admin)):
    if not CHAT_DEPS:
        return {"installed": False, "chunks": 0}
    try:
        count = _chroma_collection().count()
        return {"installed": True, "chunks": count}
    except Exception:
        return {"installed": True, "chunks": 0}

@app.get("/api/admin/chat/resume-text")
def get_resume_text(_: str = Depends(verify_admin)):
    """Return raw extracted text so admin can verify extraction quality."""
    cfg = load_config()
    url = cfg.get("resume", {}).get("file", "")
    if not url:
        raise HTTPException(400, "No resume uploaded.")
    pdf_path = DATA_DIR / url.replace("/uploads/", "", 1)
    if not pdf_path.exists():
        raise HTTPException(400, "Resume file not found.")
    try:
        text = _extract_pdf_text(pdf_path)
        chunks = _split_text(text)
        return {"text": text, "chunks": chunks, "chunk_count": len(chunks)}
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/admin/chat/ingest")
async def ingest_chat(_: str = Depends(verify_admin)):
    _check_chat_deps()
    count = ingest_resume_to_chroma()
    if count == 0:
        raise HTTPException(400, "No resume PDF found. Upload one first.")
    return {"chunks": count}

@app.post("/api/chat")
async def chat_endpoint(data: dict):
    _check_chat_deps()
    question     = data.get("question", "").strip()
    chat_summary = data.get("chat_summary", "")
    if not question:
        raise HTTPException(400, "Question required")

    # Use saved full resume text — avoids retrieval gaps from chunking/embedding mismatches
    resume_text_file = DATA_DIR / "resume" / "resume_text.txt"
    if not resume_text_file.exists():
        # Try to extract and save now
        auto = ingest_resume_to_chroma()
        if auto == 0 or not resume_text_file.exists():
            return {
                "answer":  "No resume has been uploaded yet — please upload one in the admin panel.",
                "summary": "",
            }
    context = resume_text_file.read_text(encoding="utf-8").strip()
    if not context:
        return {"answer": "The resume appears to be empty.", "summary": chat_summary}

    cfg  = load_config()
    name = cfg.get("hero", {}).get("title", "the candidate")

    system = (
        f"You are a professional assistant presenting {name}'s background to recruiters and employers. "
        "Answer using ONLY the resume excerpts provided — never invent or assume details not present. "
        "If the context truly does not contain the answer, say 'I don't know.' "
        "Format your answers in clean Markdown: use **bold** for labels, bullet points (- ) for lists, "
        "and short paragraphs. Be thorough but concise. Do not add disclaimers."
    )
    history = f"Previous conversation summary:\n{chat_summary}\n\n" if chat_summary else ""
    user_msg = f"{history}Resume excerpts:\n{context}\n\nQuestion: {question}"

    try:
        answer = _llm_chat([
            {"role": "system", "content": system},
            {"role": "user",   "content": user_msg},
        ])
    except Exception as e:
        err = str(e).lower()
        if "connection" in err or "connect" in err or "refused" in err:
            raise HTTPException(503, "Ollama is not running. Start it with: ollama serve")
        raise HTTPException(503, f"LLM error: {e}")

    try:
        sum_prompt = (
            f"Summarize this Q&A in 2 sentences for conversation memory:\n"
            f"{'Previous: ' + chat_summary + chr(10) if chat_summary else ''}"
            f"Q: {question}\nA: {answer}"
        )
        new_summary = _llm_chat([{"role": "user", "content": sum_prompt}])
    except Exception:
        new_summary = chat_summary

    return {"answer": answer, "summary": new_summary}

# ── SMTP admin ─────────────────────────────────────────────────────────────────
@app.get("/api/admin/smtp")
def get_smtp(_: str = Depends(verify_admin)):
    cfg = load_smtp()
    return {**cfg, "pass": "••••••••" if cfg.get("pass") else ""}

@app.post("/api/admin/smtp")
async def update_smtp(data: dict, _: str = Depends(verify_admin)):
    cfg = load_smtp()
    cfg["host"] = data.get("host", cfg["host"])
    cfg["port"] = int(data.get("port", cfg["port"]))
    cfg["user"] = data.get("user", cfg["user"])
    if data.get("pass") and data["pass"] != "••••••••":
        cfg["pass"] = data["pass"]
    save_smtp(cfg)
    return {**cfg, "pass": "••••••••" if cfg.get("pass") else ""}

# ── Contact form send ──────────────────────────────────────────────────────────
@app.get("/api/contact/available")
def contact_available():
    if os.environ.get("GOOGLE_SCRIPT_URL"):
        return {"available": True}
    # HF Spaces blocks outbound SMTP (port 587 is firewalled) — disable form
    if os.environ.get("SPACE_ID"):
        return {"available": False}
    saved = load_smtp()
    smtp_user = os.environ.get("SMTP_USER") or saved.get("user", "")
    smtp_pass = os.environ.get("SMTP_PASS") or saved.get("pass", "")
    return {"available": bool(smtp_user and smtp_pass)}

@app.post("/api/contact/send")
async def send_contact(data: dict):
    # env vars take priority, fall back to saved config
    smtp_host = os.environ.get("SMTP_HOST") or ""
    smtp_port_env = os.environ.get("SMTP_PORT") or ""
    smtp_user = os.environ.get("SMTP_USER") or ""
    smtp_pass = os.environ.get("SMTP_PASS") or ""

    if not smtp_user or not smtp_pass:
        saved = load_smtp()
        smtp_host = smtp_host or saved.get("host", "smtp.gmail.com")
        smtp_port_env = smtp_port_env or str(saved.get("port", 587))
        smtp_user = smtp_user or saved.get("user", "")
        smtp_pass = smtp_pass or saved.get("pass", "")

    smtp_port = int(smtp_port_env) if smtp_port_env else 587

    if not smtp_user or not smtp_pass:
        raise HTTPException(503, "Email not configured. Set SMTP credentials in Admin → Contact.")

    cfg = load_config()
    to_email = cfg.get("contact", {}).get("email") or smtp_user

    subject = f"Portfolio Contact from {data.get('name', 'Unknown')}"
    body = (
        f"Name: {data.get('name')}\n"
        f"Email: {data.get('email')}\n\n"
        f"{data.get('message')}"
    )

    # Google Apps Script relay — 100% Google, no third parties
    # Set GOOGLE_SCRIPT_URL secret in HF Spaces (see _Instructions.md)
    script_url = os.environ.get("GOOGLE_SCRIPT_URL")
    if script_url:
        import urllib.request as _ur
        payload = json.dumps({
            "subject": subject,
            "body":    body,
            "replyTo": data.get("email", ""),
        }).encode()
        req = _ur.Request(
            script_url,
            data=payload,
            headers={"Content-Type": "application/json"},
        )
        try:
            _ur.urlopen(req, timeout=15)
            return {"ok": True}
        except Exception as e:
            raise HTTPException(500, f"Email error: {e}")

    # SMTP fallback (works locally)
    if not smtp_user or not smtp_pass:
        raise HTTPException(503, "Email not configured. Set GOOGLE_SCRIPT_URL secret in HF Spaces.")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"] = smtp_user
    msg["To"] = to_email
    msg["Reply-To"] = data.get("email", smtp_user)
    msg.attach(MIMEText(body, "plain"))

    try:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=15) as server:
            server.ehlo()
            server.starttls()
            server.ehlo()
            server.login(smtp_user, smtp_pass)
            server.sendmail(smtp_user, to_email, msg.as_string())
    except smtplib.SMTPAuthenticationError:
        raise HTTPException(500, "SMTP authentication failed. Check your email and app password.")
    except (TimeoutError, OSError) as e:
        raise HTTPException(500, f"Cannot reach mail server ({smtp_host}:{smtp_port}). HuggingFace Spaces blocks outbound SMTP — set GOOGLE_SCRIPT_URL in Space secrets to enable email.")
    except Exception as e:
        raise HTTPException(500, f"Failed to send email: {e}")

    return {"ok": True}

# ── Contact ────────────────────────────────────────────────────────────────────
@app.post("/api/admin/contact")
async def update_contact(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["contact"].update(data)
    save_config(cfg)
    return cfg["contact"]

# ── Showcase Images / Slideshow ────────────────────────────────────────────────
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp"}

def _sorted_showcase_images() -> list:
    return sorted(
        [f for f in SHOWCASE_IMAGES_DIR.iterdir() if f.suffix.lower() in IMAGE_EXTS],
        key=lambda f: f.name
    )

@app.get("/api/admin/showcase/images")
def list_showcase_images(_: str = Depends(verify_admin)):
    images = [
        {"filename": f.name, "url": f"/uploads/showcase_images/{f.name}"}
        for f in _sorted_showcase_images()
    ]
    slides = sorted(SHOWCASE_IMAGES_DIR.glob("slideshow_*.mp4"), key=lambda f: f.stat().st_mtime, reverse=True)
    slideshow_url = f"/uploads/showcase_images/{slides[0].name}" if slides else (
        f"/uploads/showcase_images/slideshow.mp4" if SHOWCASE_SLIDESHOW.exists() else None
    )
    return {"images": images, "slideshow_url": slideshow_url}

@app.post("/api/admin/showcase/images")
async def upload_showcase_image(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    ext = Path(file.filename).suffix.lower()
    if ext not in IMAGE_EXTS:
        raise HTTPException(400, "Only JPG, PNG, WEBP images allowed")
    dest = SHOWCASE_IMAGES_DIR / file.filename
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    return {"filename": file.filename, "url": f"/uploads/showcase_images/{file.filename}"}

@app.delete("/api/admin/showcase/images/{filename}")
def delete_showcase_image(filename: str, _: str = Depends(verify_admin)):
    (SHOWCASE_IMAGES_DIR / filename).unlink(missing_ok=True)
    return {"ok": True}

MUSIC_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac"}

@app.get("/api/admin/showcase/music")
def list_showcase_music(_: str = Depends(verify_admin)):
    tracks = sorted(
        [f for f in SHOWCASE_MUSIC_DIR.iterdir() if f.suffix.lower() in MUSIC_EXTS],
        key=lambda f: f.name,
    )
    return [{"filename": f.name, "url": f"/uploads/showcase_music/{f.name}"} for f in tracks]

@app.post("/api/admin/showcase/music")
async def upload_showcase_music(file: UploadFile = File(...), _: str = Depends(verify_admin)):
    ext = Path(file.filename).suffix.lower()
    if ext not in MUSIC_EXTS:
        raise HTTPException(400, "Only MP3, WAV, M4A, OGG, FLAC allowed")
    dest = SHOWCASE_MUSIC_DIR / file.filename
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    return {"filename": file.filename, "url": f"/uploads/showcase_music/{file.filename}"}

@app.delete("/api/admin/showcase/music/{filename}")
def delete_showcase_music(filename: str, _: str = Depends(verify_admin)):
    (SHOWCASE_MUSIC_DIR / filename).unlink(missing_ok=True)
    return {"ok": True}

XFADE_TRANSITIONS = [
    "fade", "wipeleft", "wiperight", "wipeup", "wipedown",
    "slideleft", "slideright", "slideup", "slidedown",
    "dissolve", "fadeblack", "fadewhite",
    "horzopen", "horzclose", "vertopen", "vertclose",
    "radial", "circleopen", "circleclose",
]

@app.post("/api/admin/showcase/images/generate")
def generate_slideshow(
    duration: int = Query(default=8),
    music: str = Query(default=""),
    files: list[str] = Query(default=[]),
    _: str = Depends(verify_admin),
):
    import subprocess, random
    # Clean up stale temp files from any previous failed run
    for stale in SHOWCASE_IMAGES_DIR.glob("_clip*.mp4"):
        stale.unlink(missing_ok=True)
    for stale in SHOWCASE_IMAGES_DIR.glob("_video.mp4"):
        stale.unlink(missing_ok=True)

    all_images = _sorted_showcase_images()
    if files:
        name_map = {f.name: f for f in all_images}
        images = [name_map[fn] for fn in files if fn in name_map]
    else:
        images = all_images
    if not images:
        raise HTTPException(400, "No images uploaded yet.")

    import time as _time
    slide_name = f"slideshow_{int(_time.time())}.mp4"
    out = str(SHOWCASE_IMAGES_DIR / slide_name)
    T = 0.8  # transition duration in seconds
    SCALE = (
        "scale=1920:1080:force_original_aspect_ratio=increase,"
        "crop=1920:1080,"
        "format=yuv420p,setpts=PTS-STARTPTS"
    )

    n = len(images)
    # Exact video duration so audio trim is precise
    total_dur = round(n * duration - (n - 1) * T if n > 1 else duration, 3)

    music_path = (SHOWCASE_MUSIC_DIR / music) if music else None
    has_music  = bool(music_path and music_path.exists())

    # ── Step 1: render each image to a tmp clip (simple -vf, no filter_complex) ──
    # This avoids EINVAL from xfade on raw -loop image streams in older ffmpeg.
    clips = []
    for i, img in enumerate(images):
        tmp = str(SHOWCASE_IMAGES_DIR / f"_clip{i}.mp4")
        cmd = [
            "ffmpeg", "-y",
            "-loop", "1", "-t", str(duration), "-i", str(img),
            "-vf", f"fps=24,{SCALE}",
            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
            "-pix_fmt", "yuv420p", "-r", "24",
            "-movflags", "+faststart",
            tmp,
        ]
        r = subprocess.run(cmd, capture_output=True, text=True)
        if r.returncode != 0:
            raise HTTPException(500, f"ffmpeg clip error: {r.stderr[-400:]}")
        if not Path(tmp).exists() or Path(tmp).stat().st_size < 1000:
            raise HTTPException(500, f"ffmpeg clip {i} produced invalid output")
        clips.append(tmp)

    # ── Step 2: xfade clips into video (or single-clip pass-through) ──
    if n == 1:
        tmp_video = clips[0]
    else:
        inputs_v = []
        for c in clips:
            inputs_v += ["-i", c]
        vf = []
        prev = "0:v"
        for i in range(1, n):
            transition = random.choice(XFADE_TRANSITIONS)
            offset = round(i * duration - i * T, 3)
            lbl = "vout" if i == n - 1 else f"x{i}"
            vf.append(f"[{prev}][{i}:v]xfade=transition={transition}:duration={T}:offset={offset}[{lbl}]")
            prev = lbl
        tmp_video = str(SHOWCASE_IMAGES_DIR / "_video.mp4")
        cmd = ["ffmpeg", "-y"] + inputs_v + [
            "-filter_complex", ";".join(vf),
            "-map", "[vout]",
            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
            "-r", "24", "-t", str(total_dur),
            tmp_video,
        ]
        r = subprocess.run(cmd, capture_output=True, text=True)
        if r.returncode != 0:
            raise HTTPException(500, f"ffmpeg xfade error: {r.stderr[-400:]}")

    # ── Step 3: merge audio ──
    if has_music:
        fade_dur   = min(5.0, total_dur / 2)          # 5 s fade, or half the video if very short
        fade_start = round(total_dur - fade_dur, 3)
        cmd = [
            "ffmpeg", "-y",
            "-i", tmp_video,
            "-stream_loop", "-1", "-t", str(total_dur), "-i", str(music_path),
            "-filter_complex", f"[1:a]afade=type=out:start_time={fade_start}:duration={fade_dur}[aout]",
            "-map", "0:v", "-map", "[aout]",
            "-c:v", "copy",
            "-c:a", "aac", "-b:a", "192k",
            "-t", str(total_dur),
            out,
        ]
        r = subprocess.run(cmd, capture_output=True, text=True)
        if r.returncode != 0:
            raise HTTPException(500, f"ffmpeg audio error: {r.stderr[-400:]}")
    else:
        import shutil
        shutil.move(tmp_video, out)

    # ── Cleanup tmp files ──
    for c in clips:
        try: Path(c).unlink(missing_ok=True)
        except: pass
    if n > 1:
        try: Path(str(SHOWCASE_IMAGES_DIR / "_video.mp4")).unlink(missing_ok=True)
        except: pass

    return {"slideshow_url": f"/uploads/showcase_images/{slide_name}"}

# ── Showcase ───────────────────────────────────────────────────────────────────
@app.post("/api/admin/showcase")
async def update_showcase_meta(data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    cfg["showcase"].update({k: v for k, v in data.items() if k != "videos"})
    save_config(cfg)
    return cfg["showcase"]

@app.post("/api/admin/showcase/videos")
async def upload_showcase_video(
    file: UploadFile = File(...),
    title: str = Form(""),
    description: str = Form(""),
    _: str = Depends(verify_admin),
):
    vid_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix.lower()
    dest = DATA_DIR / "projects" / "showcase" / f"{vid_id}{ext}"
    dest.parent.mkdir(parents=True, exist_ok=True)
    async with aiofiles.open(dest, "wb") as f:
        await f.write(await file.read())
    url = f"/uploads/projects/showcase/{vid_id}{ext}"
    entry = {"id": vid_id, "url": url, "title": title, "description": description}
    cfg = load_config()
    cfg["showcase"]["videos"].append(entry)
    save_config(cfg)
    return entry

@app.post("/api/admin/showcase/videos/url")
async def add_showcase_video_url(data: dict, _: str = Depends(verify_admin)):
    vid_id = str(uuid.uuid4())
    entry = {
        "id": vid_id,
        "url": data.get("url", ""),
        "title": data.get("title", ""),
        "description": data.get("description", ""),
    }
    cfg = load_config()
    cfg["showcase"]["videos"].append(entry)
    save_config(cfg)
    return entry

@app.get("/api/admin/showcase/videos/sizes")
def showcase_video_sizes(_: str = Depends(verify_admin)):
    cfg = load_config()
    result = {}
    for v in cfg["showcase"]["videos"]:
        url = v["url"].split("?")[0]
        if url.startswith("/uploads/"):
            path = DATA_DIR / url.removeprefix("/uploads/")
            try:
                result[v["id"]] = path.stat().st_size
            except Exception:
                result[v["id"]] = None
        else:
            result[v["id"]] = None
    return result

@app.post("/api/admin/showcase/videos/{vid_id}/feature")
def feature_showcase_video(vid_id: str, _: str = Depends(verify_admin)):
    cfg = load_config()
    videos = cfg["showcase"]["videos"]
    idx = next((i for i, v in enumerate(videos) if v["id"] == vid_id), None)
    if idx is None:
        raise HTTPException(404, "Video not found")
    cfg["showcase"]["videos"] = [videos[idx]] + videos[:idx] + videos[idx+1:]
    save_config(cfg)
    return {"ok": True}

@app.patch("/api/admin/showcase/videos/{vid_id}")
def update_showcase_video(vid_id: str, data: dict, _: str = Depends(verify_admin)):
    cfg = load_config()
    videos = cfg["showcase"]["videos"]
    video = next((v for v in videos if v["id"] == vid_id), None)
    if video is None:
        raise HTTPException(404, "Video not found")
    if "title" in data:
        video["title"] = data["title"]
    if "description" in data:
        video["description"] = data["description"]
    save_config(cfg)
    return {"ok": True}

@app.delete("/api/admin/showcase/videos/{vid_id}")
def delete_showcase_video(vid_id: str, _: str = Depends(verify_admin)):
    cfg = load_config()
    before = len(cfg["showcase"]["videos"])
    cfg["showcase"]["videos"] = [v for v in cfg["showcase"]["videos"] if v["id"] != vid_id]
    save_config(cfg)
    return {"ok": True}

# ── Portfolio export (PDF / DOCX) ─────────────────────────────────────────────

def _resolve_image(img_path: str) -> Path | None:
    """Resolve a /uploads/… URL to an absolute filesystem path."""
    if not img_path:
        return None
    rel = img_path.lstrip("/")
    if rel.startswith("uploads/"):
        rel = rel[len("uploads/"):]
    p = DATA_DIR / rel
    return p if p.exists() else None


def _generate_pdf(cfg: dict) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        Image as RLImage, PageBreak, KeepTogether, Table, TableStyle,
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / "portfolio.pdf"

    acc = cfg.get("theme", {}).get("accent", "#C9A84C")
    try:
        accent = colors.HexColor(acc)
    except Exception:
        accent = colors.HexColor("#C9A84C")
    dark = colors.HexColor("#1a1a1a")
    muted = colors.HexColor("#555555")

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("H1", fontSize=28, textColor=dark, spaceAfter=4, spaceBefore=0, leading=34, fontName="Helvetica-Bold", alignment=TA_CENTER)
    h2 = ParagraphStyle("H2", fontSize=11, textColor=accent, spaceAfter=2, spaceBefore=0, leading=14, fontName="Helvetica-Bold", alignment=TA_CENTER, textTransform="uppercase")
    h3 = ParagraphStyle("H3", fontSize=14, textColor=dark, spaceAfter=6, spaceBefore=12, leading=18, fontName="Helvetica-Bold")
    sec = ParagraphStyle("SEC", fontSize=10, textColor=accent, spaceAfter=4, spaceBefore=16, leading=13, fontName="Helvetica-Bold", textTransform="uppercase")
    body = ParagraphStyle("BODY", fontSize=10, textColor=dark, spaceAfter=6, spaceBefore=0, leading=15, fontName="Helvetica", alignment=TA_JUSTIFY)
    small = ParagraphStyle("SMALL", fontSize=9, textColor=muted, spaceAfter=4, leading=13, fontName="Helvetica")
    bullet_style = ParagraphStyle("BUL", fontSize=10, textColor=dark, spaceAfter=2, leading=14, fontName="Helvetica", leftIndent=12, bulletIndent=0)

    W, H = A4
    margin = 20 * mm
    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            leftMargin=margin, rightMargin=margin,
                            topMargin=margin, bottomMargin=margin,
                            title=f"{cfg.get('hero', {}).get('title', 'Portfolio')} – Portfolio")

    story = []
    hero = cfg.get("hero", {})
    about = cfg.get("about", {})
    projects = cfg.get("projects", [])

    # ── Cover ──
    story.append(Spacer(1, 30 * mm))
    profile_path = _resolve_image(about.get("photo", ""))
    if profile_path:
        try:
            img = RLImage(str(profile_path), width=40 * mm, height=40 * mm)
            img.hAlign = "CENTER"
            story.append(img)
            story.append(Spacer(1, 6 * mm))
        except Exception:
            pass
    story.append(Paragraph(hero.get("title", "Portfolio"), h1))
    story.append(Paragraph(hero.get("subtitle", ""), h2))
    story.append(Spacer(1, 4 * mm))
    story.append(HRFlowable(width="60%", thickness=1, color=accent, hAlign="CENTER"))
    story.append(Spacer(1, 4 * mm))
    if hero.get("tagline"):
        story.append(Paragraph(hero["tagline"], ParagraphStyle("TAG", fontSize=11, textColor=muted, alignment=TA_CENTER, leading=16, fontName="Helvetica-Oblique")))
    story.append(PageBreak())

    # ── About ──
    story.append(Paragraph("About", sec))
    story.append(HRFlowable(width="100%", thickness=0.5, color=accent))
    story.append(Spacer(1, 3 * mm))
    if about.get("bio"):
        story.append(Paragraph(about["bio"], body))
    if about.get("philosophy"):
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph(about["philosophy"], body))

    # ── Skills ──
    skills = about.get("skills", [])
    if skills:
        story.append(Paragraph("Skills", sec))
        story.append(HRFlowable(width="100%", thickness=0.5, color=accent))
        story.append(Spacer(1, 2 * mm))
        cols = 3
        rows = [skills[i:i + cols] for i in range(0, len(skills), cols)]
        tdata = [[Paragraph(f"• {s}", bullet_style) for s in row] + [Paragraph("", bullet_style)] * (cols - len(row)) for row in rows]
        col_w = (W - 2 * margin) / cols
        t = Table(tdata, colWidths=[col_w] * cols)
        t.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 2)]))
        story.append(t)

    # ── Education ──
    education = about.get("education", [])
    if education:
        story.append(Paragraph("Education", sec))
        story.append(HRFlowable(width="100%", thickness=0.5, color=accent))
        story.append(Spacer(1, 2 * mm))
        for ed in education:
            story.append(Paragraph(f"<b>{ed.get('degree', '')}</b> — {ed.get('institution', '')}", body))
            if ed.get("year"):
                story.append(Paragraph(ed["year"], small))

    # ── Honors ──
    honors = about.get("honors", [])
    if honors:
        story.append(Paragraph("Honors & Awards", sec))
        story.append(HRFlowable(width="100%", thickness=0.5, color=accent))
        story.append(Spacer(1, 2 * mm))
        for h in honors:
            story.append(Paragraph(f"<b>{h.get('title', '')}</b> ({h.get('year', '')})", body))
            if h.get("description"):
                story.append(Paragraph(h["description"], small))

    # ── Publications ──
    pubs = about.get("publications", [])
    if pubs:
        story.append(Paragraph("Publications", sec))
        story.append(HRFlowable(width="100%", thickness=0.5, color=accent))
        story.append(Spacer(1, 2 * mm))
        for p in pubs:
            story.append(Paragraph(f"<b>{p.get('title', '')}</b> — {p.get('venue', '')} ({p.get('year', '')})", body))

    story.append(PageBreak())

    # ── Projects ──
    if projects:
        story.append(Paragraph("Selected Projects", sec))
        story.append(HRFlowable(width="100%", thickness=0.5, color=accent))
        story.append(Spacer(1, 3 * mm))
        usable_w = W - 2 * margin
        img_w = usable_w
        img_h = img_w * 0.56  # 16:9ish

        for proj in projects:
            block = []
            proj_title = proj.get("name") or proj.get("title", "")
            company = proj.get("title", "") if proj.get("name") else ""
            block.append(Paragraph(proj_title, h3))
            meta_parts = []
            if company:
                meta_parts.append(f"<b>Company:</b> {company}")
            if proj.get("category"):
                meta_parts.append(f"<b>Category:</b> {proj['category']}")
            if proj.get("year"):
                meta_parts.append(f"<b>Year:</b> {proj['year']}")
            if proj.get("client"):
                meta_parts.append(f"<b>Client:</b> {proj['client']}")
            if proj.get("location"):
                meta_parts.append(f"<b>Location:</b> {proj['location']}")
            if meta_parts:
                block.append(Paragraph("  |  ".join(meta_parts), small))
            if proj.get("description"):
                block.append(Paragraph(proj["description"].replace("\r\n", "\n").replace("\n", "<br/>"), body))
            # Images — show up to 3
            images = proj.get("images", [])
            if not images and proj.get("cover"):
                images = [proj["cover"]]
            shown = 0
            for img_url in images[:3]:
                img_path = _resolve_image(img_url)
                if img_path:
                    try:
                        rl_img = RLImage(str(img_path), width=img_w, height=img_h)
                        rl_img.hAlign = "LEFT"
                        block.append(Spacer(1, 2 * mm))
                        block.append(rl_img)
                        shown += 1
                    except Exception:
                        pass
            block.append(Spacer(1, 6 * mm))
            block.append(HRFlowable(width="100%", thickness=0.3, color=colors.HexColor("#dddddd")))
            story.append(KeepTogether(block[:4]))  # keep header+meta+desc together
            for item in block[4:]:
                story.append(item)

    doc.build(story)
    return out


def _generate_docx(cfg: dict) -> Path:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / "portfolio.docx"

    def hex_to_rgb(h: str):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    acc_hex = cfg.get("theme", {}).get("accent", "C9A84C").lstrip("#")
    acc_r, acc_g, acc_b = hex_to_rgb(acc_hex)
    accent_color = RGBColor(acc_r, acc_g, acc_b)
    dark_color = RGBColor(26, 26, 26)
    muted_color = RGBColor(85, 85, 85)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    def add_heading(text, level=1, color=None, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        sizes = {1: 26, 2: 14, 3: 11}
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = color or dark_color
        return p

    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = accent_color
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), acc_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_body(text, italic=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = color or dark_color
        run.italic = italic
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_meta(parts: list[str]):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        for i, part in enumerate(parts):
            if ":" in part:
                label, val = part.split(":", 1)
                run = p.add_run(label + ":")
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = muted_color
                run2 = p.add_run(val)
                run2.font.size = Pt(9)
                run2.font.color.rgb = muted_color
            else:
                run = p.add_run(part)
                run.font.size = Pt(9)
                run.font.color.rgb = muted_color
            if i < len(parts) - 1:
                sep = p.add_run("  |  ")
                sep.font.size = Pt(9)
                sep.font.color.rgb = muted_color
        return p

    hero = cfg.get("hero", {})
    about = cfg.get("about", {})
    projects = cfg.get("projects", [])

    # ── Cover ──
    profile_path = _resolve_image(about.get("photo", ""))
    if profile_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(profile_path), width=Inches(1.5))
        except Exception:
            pass

    add_heading(hero.get("title", "Portfolio"), level=1, center=True)
    sub_p = add_heading(hero.get("subtitle", ""), level=3, color=accent_color, center=True)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if hero.get("tagline"):
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(hero["tagline"])
        tr.italic = True
        tr.font.size = Pt(11)
        tr.font.color.rgb = muted_color
    doc.add_page_break()

    # ── About ──
    add_section_heading("About")
    if about.get("bio"):
        add_body(about["bio"])
    if about.get("philosophy"):
        doc.add_paragraph()
        add_body(about["philosophy"])

    # ── Skills ──
    skills = about.get("skills", [])
    if skills:
        add_section_heading("Skills")
        cols = 3
        rows = [skills[i:i + cols] for i in range(0, len(skills), cols)]
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci, skill in enumerate(row):
                cell = table.cell(ri, ci)
                cell.text = f"• {skill}"
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # ── Education ──
    education = about.get("education", [])
    if education:
        add_section_heading("Education")
        for ed in education:
            p = doc.add_paragraph()
            r1 = p.add_run(f"{ed.get('degree', '')}")
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = dark_color
            r2 = p.add_run(f" — {ed.get('institution', '')}")
            r2.font.size = Pt(10)
            r2.font.color.rgb = dark_color
            if ed.get("year"):
                yp = doc.add_paragraph(ed["year"])
                yp.runs[0].font.size = Pt(9)
                yp.runs[0].font.color.rgb = muted_color
                yp.paragraph_format.space_after = Pt(2)

    # ── Honors ──
    honors = about.get("honors", [])
    if honors:
        add_section_heading("Honors & Awards")
        for h in honors:
            p = doc.add_paragraph()
            r1 = p.add_run(h.get("title", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = dark_color
            r2 = p.add_run(f" ({h.get('year', '')})")
            r2.font.size = Pt(9)
            r2.font.color.rgb = muted_color
            if h.get("description"):
                add_body(h["description"], color=muted_color)

    # ── Publications ──
    pubs = about.get("publications", [])
    if pubs:
        add_section_heading("Publications")
        for pub in pubs:
            p = doc.add_paragraph()
            r1 = p.add_run(pub.get("title", ""))
            r1.bold = True
            r1.font.size = Pt(10)
            r1.font.color.rgb = dark_color
            r2 = p.add_run(f" — {pub.get('venue', '')} ({pub.get('year', '')})")
            r2.font.size = Pt(9)
            r2.font.color.rgb = muted_color

    doc.add_page_break()

    # ── Projects ──
    if projects:
        add_section_heading("Selected Projects")
        for proj in projects:
            proj_title = proj.get("name") or proj.get("title", "")
            company = proj.get("title", "") if proj.get("name") else ""
            add_heading(proj_title, level=2)
            meta_parts = []
            if company:
                meta_parts.append(f"Company: {company}")
            if proj.get("category"):
                meta_parts.append(f"Category: {proj['category']}")
            if proj.get("year"):
                meta_parts.append(f"Year: {proj['year']}")
            if proj.get("client"):
                meta_parts.append(f"Client: {proj['client']}")
            if proj.get("location"):
                meta_parts.append(f"Location: {proj['location']}")
            if meta_parts:
                add_meta(meta_parts)
            if proj.get("description"):
                clean_desc = proj["description"].replace("\r\n", "\n").strip()
                add_body(clean_desc)
            images = proj.get("images", [])
            if not images and proj.get("cover"):
                images = [proj["cover"]]
            for img_url in images[:3]:
                img_path = _resolve_image(img_url)
                if img_path:
                    try:
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(str(img_path), width=Inches(5.5))
                    except Exception:
                        pass
            doc.add_paragraph()

    doc.save(str(out))
    return out


@app.get("/api/portfolio/download")
def download_portfolio(fmt: str = Query("pdf", regex="^(pdf|docx)$")):
    """Generate and return electronic_portfolio.pdf or electronic_portfolio.docx."""
    cfg = load_config()
    if fmt == "docx":
        out = _generate_docx(cfg)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "electronic_portfolio.docx"
    else:
        out = _generate_pdf(cfg)
        media = "application/pdf"
        filename = "electronic_portfolio.pdf"
    return FileResponse(str(out), media_type=media, filename=filename)


# ── Serve media/ folder (pre-loaded portfolio PDF, etc.) ──────────────────────
if MEDIA_DIR.exists():
    app.mount("/media", StaticFiles(directory=str(MEDIA_DIR)), name="media")

# ── Serve uploaded files ───────────────────────────────────────────────────────
app.mount("/uploads", StaticFiles(directory=str(DATA_DIR)), name="uploads")

# ── Serve React SPA (must be last) ────────────────────────────────────────────
if STATIC_DIR.exists():
    app.mount("/assets", StaticFiles(directory=str(STATIC_DIR / "assets")), name="assets")

    @app.api_route("/{full_path:path}", methods=["GET", "HEAD"])
    def serve_spa(full_path: str):
        index = STATIC_DIR / "index.html"
        return FileResponse(str(index))
